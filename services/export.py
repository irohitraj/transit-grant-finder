"""
Export the approved narrative draft to Markdown or DOCX.

Neither function writes files to disk — both return bytes that Streamlit's
st.download_button can serve directly.
"""

import io
from datetime import datetime
from typing import Dict

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _build_header(grant_name: str, agency_name: str) -> str:
    date_str = datetime.now().strftime("%B %d, %Y")
    return f"Grant: {grant_name}\nAgency: {agency_name}\nPrepared: {date_str}"


def export_markdown(
    narrative: str,
    grant_name: str,
    agency_name: str,
) -> bytes:
    """Return a UTF-8 encoded Markdown document as bytes.

    Args:
        narrative: The drafted narrative text (may contain placeholders).
        grant_name: Name of the selected grant program.
        agency_name: Agency name from the profile (or placeholder string).

    Returns:
        bytes — suitable for st.download_button(data=...).
    """
    header = _build_header(grant_name, agency_name)
    disclaimer = (
        "> **Note:** This draft was generated to assist in grant writing. "
        "All `[AGENCY TO PROVIDE: ...]` placeholders must be completed by the agency "
        "before submission. Do not submit this document without review."
    )
    content = "\n\n".join(
        [
            f"# Grant Narrative Draft",
            header,
            "---",
            disclaimer,
            "---",
            "## Project Narrative",
            narrative,
        ]
    )
    return content.encode("utf-8")


def export_docx(
    narrative: str,
    grant_name: str,
    agency_name: str,
) -> bytes:
    """Return a DOCX document as bytes.

    Args:
        narrative: The drafted narrative text (may contain placeholders).
        grant_name: Name of the selected grant program.
        agency_name: Agency name from the profile (or placeholder string).

    Returns:
        bytes — suitable for st.download_button(data=...).
    """
    doc = Document()

    # Title
    title = doc.add_heading("Grant Narrative Draft", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata block
    date_str = datetime.now().strftime("%B %d, %Y")
    meta_para = doc.add_paragraph()
    meta_para.add_run("Grant: ").bold = True
    meta_para.add_run(grant_name)
    meta_para.add_run("\nAgency: ").bold = True
    meta_para.add_run(agency_name)
    meta_para.add_run("\nPrepared: ").bold = True
    meta_para.add_run(date_str)

    doc.add_paragraph()  # spacer

    # Disclaimer
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "Note: This draft was generated to assist in grant writing. "
        "All [AGENCY TO PROVIDE: ...] placeholders must be completed by the agency "
        "before submission. Do not submit this document without review."
    )
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Section heading
    doc.add_heading("Project Narrative", level=2)

    # Narrative body — split on double newlines to preserve paragraph breaks
    paragraphs = narrative.strip().split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
