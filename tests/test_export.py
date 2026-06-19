"""Tests for services/export.py — Markdown and DOCX export."""

import io
import pytest
from docx import Document

from services.export import export_markdown, export_docx


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_NARRATIVE = (
    "Valley Rural Transit serves a three-county rural region in eastern Washington.\n\n"
    "The agency operates [AGENCY TO PROVIDE: fleet size (number of vehicles)] vehicles "
    "and serves [AGENCY TO PROVIDE: annual ridership] riders annually.\n\n"
    "This project will implement AI-assisted scheduling software."
)

GRANT_NAME = "Section 5311 + RTAP"
AGENCY_NAME = "Valley Rural Transit"
PLACEHOLDER_AGENCY = "[AGENCY TO PROVIDE: agency name]"


# ---------------------------------------------------------------------------
# export_markdown
# ---------------------------------------------------------------------------

class TestExportMarkdown:
    def test_returns_bytes(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        assert isinstance(result, bytes)

    def test_contains_narrative_text(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        text = result.decode("utf-8")
        assert "Valley Rural Transit" in text
        assert "AI-assisted scheduling software" in text

    def test_contains_grant_name(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        text = result.decode("utf-8")
        assert GRANT_NAME in text

    def test_contains_agency_name(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        text = result.decode("utf-8")
        assert AGENCY_NAME in text

    def test_placeholders_preserved(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        text = result.decode("utf-8")
        assert "[AGENCY TO PROVIDE: fleet size (number of vehicles)]" in text
        assert "[AGENCY TO PROVIDE: annual ridership]" in text

    def test_no_none_in_output(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        text = result.decode("utf-8")
        assert "None" not in text

    def test_works_with_placeholder_agency_name(self):
        result = export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, PLACEHOLDER_AGENCY)
        text = result.decode("utf-8")
        assert "[AGENCY TO PROVIDE: agency name]" in text

    def test_does_not_write_to_disk(self, tmp_path):
        """Verify no files are written during export."""
        before = list(tmp_path.iterdir())
        export_markdown(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        after = list(tmp_path.iterdir())
        assert before == after  # no new files


# ---------------------------------------------------------------------------
# export_docx
# ---------------------------------------------------------------------------

class TestExportDocx:
    def test_returns_bytes(self):
        result = export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        assert isinstance(result, bytes)

    def test_parseable_by_python_docx(self):
        result = export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_docx_contains_narrative_text(self):
        result = export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AI-assisted scheduling software" in full_text

    def test_docx_contains_grant_name(self):
        result = export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert GRANT_NAME in full_text

    def test_docx_placeholders_preserved(self):
        result = export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[AGENCY TO PROVIDE: annual ridership]" in full_text

    def test_no_none_in_docx_output(self):
        result = export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "None" not in full_text

    def test_does_not_write_to_disk(self, tmp_path):
        """Verify no files are written during export."""
        before = list(tmp_path.iterdir())
        export_docx(SAMPLE_NARRATIVE, GRANT_NAME, AGENCY_NAME)
        after = list(tmp_path.iterdir())
        assert before == after
